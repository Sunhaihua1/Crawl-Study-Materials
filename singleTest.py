from selenium import webdriver
import re
import docx
from docx import Document
import time


def writeWord(list):
    document = Document()
    for x in range(len(list)):
        document.add_paragraph(list[x]["testLine"])
        document.add_paragraph(list[x]["A"])
        document.add_paragraph(list[x]["B"])
        document.add_paragraph(list[x]["C"])
        document.add_paragraph(list[x]["D"])
        document.add_paragraph(list[x]["answer"])
        document.add_paragraph()

    document.save('1.docx')



def getData(brower):
    allInfos = []

    for i in range(89):
        metaInfos = {}

        time.sleep(1)
        outAnswer = brower.find_element_by_xpath(
            '//*[@id="questionRender"]/div[2]/div[2]/div[3]/div[2]/label[1]/div[1]/input')
        if i != 11:
            outAnswer.click()


        testLineTitle = brower.find_element_by_xpath(
            '//*[ @ id = "questionRender"] / div[2] / div[2] / div[3] / div[1]/div[1]'
        )
        testLineMes = brower.find_element_by_xpath(
            '//*[ @ id = "questionRender"] / div[2] / div[2] / div[3] / div[1]/div[2]'
        )
        testLine = str(testLineTitle.text) + str(testLineMes.text)
        metaInfos['testLine'] = testLine
       # print(testLine)

        for n in range(1, 5):
            choiceTitle = brower.find_element_by_xpath(
                '//*[@id="questionRender"]/div[2]/div[2]/div[3]/div[2]/label[' + str(n) + ']/div[1]'
            )
            choiceMes = brower.find_element_by_xpath(
                '//*[@id="questionRender"]/div[2]/div[2]/div[3]/div[2]/label[' + str(n) + ']/div[2]'
            )
            choice = str(choiceTitle.text) + str(choiceMes.text)
            if n==1:
                metaInfos['A'] = choice
            elif n==2:
                metaInfos['B'] = choice
            elif n==3:
                metaInfos['C'] = choice
            else:
                metaInfos['D'] = choice
                #print(metaInfos['D'])
        if i != 11:
            answerTitle = brower.find_element_by_xpath(
                '//*[@id="questionRender"]/div[3]/div[2]/div[1]/div[2]/div/div[1]'
            )
            answerMes = brower.find_element_by_xpath(
                '//*[@id="questionRender"]/div[3]/div[2]/div[1]/div[2]/div/div[2]'
            )
            answer = str(answerTitle.text) + str(answerMes.text)
            print(answer)
            metaInfos['answer'] = answer
            allInfos.append(metaInfos)
        else:
            metaInfos['answer'] = "参考答案：A"


        # print()
        if i == 88:
            break
        nextPage = brower.find_element_by_xpath('//*[@id="sectionRender"]/div[1]/div[2]/a[' + str(i+2) + ']')
        nextPage.click()

    #print(allInfos)

    return allInfos

brower = webdriver.Chrome()
brower.get("http://saishi.cnki.net/Exercise/Practice/ks0af38326-10ce-48a5-a2cb-84204baf254c/1")
writeWord(getData(brower))
